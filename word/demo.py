pip install matplotlib python-docx

# mindmap.py
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Step 1: Define your mind map structure
main_idea = "Healthy Lifestyle"
branches = ["Exercise", "Nutrition", "Sleep", "Hydration", "Mental Health"]

# Step 2: Draw the mind map
plt.figure(figsize=(6,6))
plt.text(0.5, 0.5, main_idea, ha='center', va='center', fontsize=14, weight='bold', color='darkgreen')

positions = [
    (0.5, 0.9),
    (0.9, 0.5),
    (0.5, 0.1),
    (0.1, 0.5),
    (0.2, 0.8)
]

for branch, pos in zip(branches, positions):
    plt.text(pos[0], pos[1], branch, ha='center', va='center', fontsize=11, color='blue')
    plt.plot([0.5, pos[0]], [0.5, pos[1]], 'k-', linewidth=1)

plt.axis('off')
plt.savefig("mindmap.png", bbox_inches='tight')
plt.close()

# Step 3: Create a Word document and insert the image
doc = Document()
doc.add_heading("Mind Map: Healthy Lifestyle", level=1)
doc.add_picture("mindmap.png", width=Inches(5))
doc.add_paragraph("This mind map shows key aspects of maintaining a healthy lifestyle.")
doc.save("mindmap.docx")

print("✅ Mind map created successfully! Check the file 'mindmap.docx'.")


